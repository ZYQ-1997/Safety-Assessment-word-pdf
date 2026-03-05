from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
import pdfplumber
import pandas as pd
import os
import json
from werkzeug.utils import secure_filename
from datetime import datetime
import zipfile
import io
import re
from typing import List, Dict, Tuple, Optional

# ---------------------------------------------------------------------------
# 配置：环境变量 + 基于 __file__ 的路径（兼容 Windows/Linux 与任意工作目录）
# ---------------------------------------------------------------------------
_BACKEND_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.dirname(_BACKEND_DIR)
FRONTEND_DIR = os.path.normpath(os.path.join(_BACKEND_DIR, "..", "frontend"))

def _safe_path(value: str, default: str) -> str:
    """路径校验：禁止 '..' 与绝对路径。"""
    if not value or ".." in value or value.startswith("/") or re.match(r"^[A-Za-z]:", value):
        return default
    return value.strip()

_PORT = os.getenv("PORT", "5000")
PORT = int(_PORT) if str(_PORT).strip().isdigit() else 5000

_DEBUG = (os.getenv("FLASK_DEBUG", "false") or "").strip().lower()
FLASK_DEBUG = _DEBUG in ("1", "true", "yes")

_upload_name = _safe_path(os.getenv("UPLOAD_FOLDER", "uploads") or "uploads", "uploads")
_output_name = _safe_path(os.getenv("OUTPUT_FOLDER", "outputs") or "outputs", "outputs")
UPLOAD_FOLDER = os.path.normpath(os.path.join(_PROJECT_ROOT, _upload_name))
OUTPUT_FOLDER = os.path.normpath(os.path.join(_PROJECT_ROOT, _output_name))

_MAX_MB = os.getenv("MAX_CONTENT_LENGTH_MB", "500")
MAX_CONTENT_LENGTH_MB = int(_MAX_MB) if str(_MAX_MB).strip().isdigit() else 500
MAX_CONTENT_LENGTH_MB = max(1, min(500, MAX_CONTENT_LENGTH_MB))
MAX_CONTENT_LENGTH = MAX_CONTENT_LENGTH_MB * 1024 * 1024

# ---------------------------------------------------------------------------

app = Flask(__name__, static_folder=FRONTEND_DIR, static_url_path="")
CORS(app, resources={r"/api/*": {"origins": "*"}})

ALLOWED_EXTENSIONS = {'pdf', 'docx'}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH
app.config['DEBUG'] = FLASK_DEBUG

# 添加全局错误处理器，确保所有错误都返回JSON格式
@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': '资源未找到', 'error_type': 'NotFound'}), 404

@app.errorhandler(500)
def internal_error(error):
    import traceback
    error_trace = traceback.format_exc()
    print(f"=" * 60)
    print(f"未捕获的服务器错误:")
    print(error_trace)
    print(f"=" * 60)
    return jsonify({
        'error': '服务器内部错误',
        'error_type': 'InternalServerError',
        'details': error_trace if app.debug else '启用debug模式查看详细信息'
    }), 500

@app.errorhandler(Exception)
def handle_exception(e):
    import traceback
    error_trace = traceback.format_exc()
    print(f"=" * 60)
    print(f"未捕获的异常:")
    print(f"异常类型: {type(e).__name__}")
    print(f"异常消息: {str(e)}")
    print(f"堆栈跟踪:\n{error_trace}")
    print(f"=" * 60)
    return jsonify({
        'error': f'服务器错误: {str(e)}',
        'error_type': type(e).__name__,
        'details': error_trace if app.debug else '启用debug模式查看详细信息'
    }), 500

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def _is_docx_by_magic(filepath: str) -> bool:
    """通过文件头魔数判断是否为 Word (.docx) 文件（docx 实为 ZIP，头为 PK）"""
    try:
        with open(filepath, "rb") as f:
            return f.read(2) == b"PK"
    except Exception:
        return False


def word_remove_non_table_content(
    docx_path: str,
    output_path: str,
    selected_table_ids: Optional[List[str]] = None,
) -> int:
    """
    原样保留 Word 中所有表格，及每个表格上面的一行文字作为表名；删除其余内容。
    表格不做任何处理，保留原文件状态。
    selected_table_ids 若提供则只保留这些表格（及各自上一行表名），否则保留全部。
    返回保留的表格数量。
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(docx_path)
    body = doc.element.body
    tbl_tag = qn("w:tbl")
    p_tag = qn("w:p")
    children = list(body)
    keep_indices = set()
    table_index = 0
    for i, child in enumerate(children):
        if child.tag == tbl_tag:
            if selected_table_ids is None or f"table_{table_index}" in selected_table_ids:
                keep_indices.add(i)
                if i > 0 and children[i - 1].tag == p_tag:
                    keep_indices.add(i - 1)
            table_index += 1
    for i in range(len(children) - 1, -1, -1):
        if i not in keep_indices:
            body.remove(children[i])
    doc.save(output_path)
    return len([i for i in keep_indices if children[i].tag == tbl_tag])

# 定义需要提取的表格（只提取表格，不提取文字）
TARGET_SECTIONS = [
    ("评价项目基本概况表", "table"),
    ('企业"两重点一重大"及剧毒化学品等辨识结果', "table"),
    ("企业与外部建、构筑物间距一览表", "table"),
    ("企业主要建（构）筑物的情况", "table"),
    ("主要产品及生产情况", "table"),
    ("产品品种及生产规模", "table"),
    ("危险化学品物料回收情况", "table"),
    ("控制系统工艺参数测量、报警、联锁保护一览表", "table"),
    ("SIS 控制系统工艺参数测量、报警、联锁保护一览表", "table"),
    ("气体浓度探测器设置一览表", "table"),
    ("重点监管工艺车间主要工艺设备一览表", "table"),
    ("其他车间主要工艺设备一览表", "table"),
    ("特种设备一览表", "table"),
    ("主要原辅材料消耗、储存一览表", "table"),
    ("甲类罐区储罐情况一览表", "table"),
    ("甲类仓库使用情况一览表", "table"),
    ("消防、应急设施配置一览表", "table"),
    ("反应风险分析数据", "table"),
    ("反应产品、中间产物、蒸馏底物等 DSC 数据", "table"),
    ("附录：危险化学品物料特性表", "table"),
]

def find_section_in_text(text: str, section_name: str) -> bool:
    """在文本中查找章节标题"""
    # 移除可能的空格和标点差异
    section_clean = section_name.replace(" ", "").replace("（", "(").replace("）", ")")
    text_clean = text.replace(" ", "").replace("（", "(").replace("）", ")")
    
    # 直接匹配
    if section_name in text or section_clean in text_clean:
        return True
    
    # 模糊匹配（移除部分标点）
    section_keywords = re.sub(r'[^\w\u4e00-\u9fff]', '', section_name)
    text_keywords = re.sub(r'[^\w\u4e00-\u9fff]', '', text)
    
    if section_keywords in text_keywords and len(section_keywords) > 5:
        return True
    
    return False

def extract_text_from_page(page) -> str:
    """从页面提取文本"""
    try:
        text = page.extract_text()
        return text if text else ""
    except:
        return ""

def extract_table_improved(page, table_settings=None) -> List:
    """改进的表格提取方法，处理复杂格式 - 优化版，确保提取所有行"""
    all_tables = []
    
    # 策略1: 使用默认设置（适合大多数标准表格）
    try:
        tables = page.extract_tables()
        if tables:
            all_tables.extend(tables)
    except:
        pass
    
    # 策略2: 使用线条检测（适合有明确边框的表格）
    try:
        tables = page.extract_tables(table_settings={
            "vertical_strategy": "lines",
            "horizontal_strategy": "lines",
            "snap_tolerance": 5,
            "join_tolerance": 3,
        })
        if tables:
            all_tables.extend(tables)
    except:
        pass
    
    # 策略3: 使用严格线条检测
    try:
        tables = page.extract_tables(table_settings={
            "vertical_strategy": "lines_strict",
            "horizontal_strategy": "lines_strict",
            "snap_tolerance": 3,
            "join_tolerance": 2,
        })
        if tables:
            all_tables.extend(tables)
    except:
        pass
    
    # 策略4: 基于文本对齐（适合无边框表格）
    try:
        tables = page.extract_tables(table_settings={
            "vertical_strategy": "text",
            "horizontal_strategy": "text",
        })
        if tables:
            all_tables.extend(tables)
    except:
        pass
    
    # 策略5: 更宽松的设置，确保不遗漏行
    try:
        tables = page.extract_tables(table_settings={
            "vertical_strategy": "lines",
            "horizontal_strategy": "lines",
            "snap_tolerance": 10,
            "join_tolerance": 5,
            "edge_tolerance": 5,
        })
        if tables:
            all_tables.extend(tables)
    except:
        pass
    
    # 去重：比较表格内容，保留最完整的（优先保留行数最多的）
    unique_tables = []
    for table in all_tables:
        if not table or len(table) < 1:  # 至少要有表头
            continue
        
        # 检查是否与已有表格重复
        is_duplicate = False
        for existing_table in unique_tables:
            if existing_table and len(existing_table) > 0:
                # 简单比较：如果第一行和第一列相同，认为是重复
                if (table[0] and existing_table[0] and 
                    table[0][:3] == existing_table[0][:3]):
                    is_duplicate = True
                    # 如果新表格更大（行数更多），替换旧的
                    if len(table) > len(existing_table):
                        unique_tables.remove(existing_table)
                        unique_tables.append(table)
                    break
        
        if not is_duplicate:
            unique_tables.append(table)
    
    # 返回行数最多的表格（通常更完整）
    if unique_tables:
        best_table = max(unique_tables, key=lambda t: len(t) if t else 0)
        # 确保表格至少有表头
        if best_table and len(best_table) > 0:
            return [best_table]
    
    return []

def clean_table_data(table: List) -> List:
    """清理表格数据，保留原始结构 - 优化版，确保不遗漏任何行，保留换行符"""
    if not table:
        return []
    
    cleaned = []
    for row in table:
        if row:
            # 清理每行的数据
            cleaned_row = []
            for cell in row:
                if cell:
                    # 保留换行符，只移除首尾空白（不包含换行符）
                    cell_str = str(cell)
                    # 移除首尾空白，但保留换行符
                    cell_str = cell_str.rstrip(' \t').lstrip(' \t')
                    # 确保换行符被保留（\n）
                    cleaned_row.append(cell_str if cell_str else None)
                else:
                    cleaned_row.append(None)
            
            # 更宽松的判断：只要行中有任何非空单元格，就保留该行
            # 不要因为某些单元格为空就删除整行
            has_content = any(cell for cell in cleaned_row if cell and str(cell).strip())
            if has_content or len(cleaned_row) > 0:  # 即使看起来是空行，也保留（可能是格式问题）
                cleaned.append(cleaned_row)
    
    return cleaned

def detect_merged_cells(table: List) -> List[Tuple[int, int, int, int]]:
    """检测表格中的合并单元格（基于内容重复判断）- 已禁用，保持字段独立"""
    # 用户要求：相同的字段不要合并，保持独立
    # 因此返回空列表，不进行任何合并
    return []

def is_table_ended(page, section_name: str) -> bool:
    """判断表格是否已经结束（基于页面内容特征）- 优化版，更早检测表格结束"""
    try:
        text = extract_text_from_page(page)
        if not text:
            return False
        
        # 检查是否出现了新的章节标题（表格结束的标志）
        for target_name, _ in TARGET_SECTIONS:
            if target_name != section_name and find_section_in_text(text, target_name):
                # 找到了其他章节标题，说明当前表格已结束
                return True
        
        # 检查是否出现了明显的章节分隔符（如"附录"、"附件"等）
        end_markers = ["附录", "附件", "参考文献", "注：", "说明：", "备注：", "表", "图"]
        for marker in end_markers:
            if marker in text:
                # 如果标记出现在页面顶部附近，可能是新章节开始
                marker_pos = text.find(marker)
                if marker_pos < len(text) * 0.4:  # 在页面前40%位置（更早检测）
                    # 进一步检查：如果标记后面跟着数字或表格相关文字，很可能是新表格
                    marker_text = text[marker_pos:marker_pos+20]
                    if any(keyword in marker_text for keyword in ["表", "一览表", "情况表", "结果表"]):
                        return True
        
        # 检查页面是否主要是文本而非表格（如果页面顶部没有表格结构，可能是新章节）
        try:
            tables = page.extract_tables()
            # 如果页面顶部（前50%区域）没有表格，且有很多文本，可能是新章节
            if not tables or len(tables) == 0:
                # 检查文本密度，如果文本很长但没表格，可能是新章节
                if len(text) > 500:  # 文本较长
                    return True
        except:
            pass
        
        return False
    except:
        return False

def clean_duplicate_cells(table: List) -> List:
    """清理破坏连贯性的多余单元格（如重复的表头、空行等）"""
    if not table or len(table) < 2:
        return table
    
    cleaned = []
    header = table[0] if table else []
    
    # 保留表头
    if header:
        cleaned.append(header)
    
    # 处理数据行
    for i, row in enumerate(table[1:], start=1):
        if not row:
            continue
        
        # 检查是否是重复的表头行（与第一行完全相同或高度相似）
        if header:
            header_similarity = calculate_header_similarity(header, row)
            if header_similarity >= 0.8:  # 80%相似度认为是重复表头
                print(f"    跳过重复表头行（第{i+1}行）")
                continue
        
        # 检查是否是空行（所有单元格都为空或只有空白）
        has_content = any(cell and str(cell).strip() for cell in row)
        if not has_content:
            continue
        
        # 检查是否是破坏连贯性的多余行（例如：只有1-2个非空单元格，且内容很短）
        non_empty_count = sum(1 for cell in row if cell and str(cell).strip())
        if non_empty_count <= 2 and len(row) > 5:
            # 检查内容长度
            max_content_len = max((len(str(cell).strip()) for cell in row if cell and str(cell).strip()), default=0)
            if max_content_len < 5:  # 内容很短，可能是多余的
                print(f"    跳过可能的多余行（第{i+1}行，非空单元格: {non_empty_count}）")
                continue
        
        cleaned.append(row)
    
    return cleaned

def has_header_row(row: List) -> bool:
    """判断一行是否是表头行"""
    if not row:
        return False
    
    # 检查非空单元格
    non_empty_cells = [str(cell).strip() for cell in row if cell and str(cell).strip()]
    if len(non_empty_cells) < 2:  # 表头至少要有2列
        return False
    
    # 表头特征：通常是文本，不全是数字
    text_count = 0
    number_count = 0
    
    for cell in non_empty_cells:
        # 尝试判断是否是数字
        try:
            float(cell.replace(',', '').replace(' ', ''))
            number_count += 1
        except:
            text_count += 1
    
    # 如果文本占比超过50%，可能是表头
    total = text_count + number_count
    if total > 0:
        return text_count / total >= 0.5
    
    return False

def is_same_table(table1: List, table2: List, strict_mode: bool = False) -> bool:
    """判断两个表格是否是同一个表格（用于跨页表格合并）- 增强版，支持无表头判断"""
    if not table1 or not table2:
        return False
    
    # 检查table2是否有表头
    table2_has_header = False
    if len(table2) > 0:
        table2_has_header = has_header_row(table2[0])
    
    # 如果table2没有表头，很可能是table1的延续
    if not table2_has_header:
        # 检查列数是否相同或接近
        if len(table1) > 0 and len(table2) > 0:
            cols1 = len([c for c in table1[0] if c])
            cols2 = len([c for c in table2[0] if c])
            
            # 如果列数相同或接近（允许±1的差异）
            if abs(cols1 - cols2) <= 1 and cols1 > 0:
                # 检查数据格式是否相似
                if len(table1) > 1 and len(table2) > 0:
                    # 比较table1的最后一行和table2的第一行格式
                    last_row1 = table1[-1] if len(table1) > 1 else []
                    first_row2 = table2[0]
                    
                    non_empty1 = sum(1 for c in last_row1 if c and str(c).strip())
                    non_empty2 = sum(1 for c in first_row2 if c and str(c).strip())
                    
                    # 如果非空单元格数量相似，很可能是同一表格
                    if abs(non_empty1 - non_empty2) <= 2:
                        return True
    
    # 方法1: 比较表头（第一行）- 最可靠的方法
    if len(table1) > 0 and len(table2) > 0:
        header1 = [str(cell).strip() if cell else "" for cell in table1[0]]
        header2 = [str(cell).strip() if cell else "" for cell in table2[0]]
        
        # 清理空列
        header1 = [h for h in header1 if h]
        header2 = [h for h in header2 if h]
        
        if len(header1) > 0 and len(header2) > 0:
            # 如果列数相同或接近
            if abs(len(header1) - len(header2)) <= 1:
                # 计算相似度
                min_len = min(len(header1), len(header2))
                match_count = sum(1 for h1, h2 in zip(header1[:min_len], header2[:min_len]) 
                                 if h1 and h2 and (h1 == h2 or h1 in h2 or h2 in h1))
                similarity = match_count / min_len if min_len > 0 else 0
                
                # 降低阈值，更宽松的判断（60%相似即可）
                if similarity >= 0.6:
                    return True
    
    # 方法2: 如果表头不匹配，检查列数是否相同（可能是表头被重复）
    if len(table1) > 0 and len(table2) > 0:
        cols1 = len([c for c in table1[0] if c])
        cols2 = len([c for c in table2[0] if c])
        
        # 如果列数相同，且表格结构相似
        if cols1 == cols2 and cols1 > 0:
            # 检查前几行数据的格式是否相似
            sample_rows = min(3, len(table1), len(table2))
            if sample_rows > 1:
                format_match = 0
                for i in range(1, sample_rows):
                    row1 = table1[i] if i < len(table1) else []
                    row2 = table2[i] if i < len(table2) else []
                    # 检查非空单元格的数量是否相似
                    non_empty1 = sum(1 for c in row1 if c and str(c).strip())
                    non_empty2 = sum(1 for c in row2 if c and str(c).strip())
                    if abs(non_empty1 - non_empty2) <= 1:
                        format_match += 1
                
                if format_match >= sample_rows - 1:
                    return True
    
    # 方法3: 检查第二页表格的第一行是否是表头重复（跨页表格常见情况）
    if len(table2) > 1:
        # 如果table2的第一行看起来像表头（与table1的表头相似）
        header2_first = [str(cell).strip() if cell else "" for cell in table2[0]]
        header1_clean = [str(cell).strip() if cell else "" for cell in table1[0]]
        
        if len(header2_first) > 0 and len(header1_clean) > 0:
            min_len = min(len(header2_first), len(header1_clean))
            match_count = sum(1 for h1, h2 in zip(header1_clean[:min_len], header2_first[:min_len])
                             if h1 and h2 and (h1 == h2 or h1 in h2 or h2 in h1))
            similarity = match_count / min_len if min_len > 0 else 0
            
            if similarity >= 0.7:
                return True
    
    return False

def calculate_header_similarity(header1: List, header2: List) -> float:
    """计算两个表头的相似度"""
    if not header1 or not header2:
        return 0.0
    
    h1 = [str(cell).strip() if cell else "" for cell in header1]
    h2 = [str(cell).strip() if cell else "" for cell in header2]
    
    h1 = [h for h in h1 if h]
    h2 = [h for h in h2 if h]
    
    if len(h1) == 0 or len(h2) == 0:
        return 0.0
    
    min_len = min(len(h1), len(h2))
    match_count = sum(1 for h1_cell, h2_cell in zip(h1[:min_len], h2[:min_len])
                     if h1_cell and h2_cell and (h1_cell == h2_cell or h1_cell in h2_cell or h2_cell in h1_cell))
    
    return match_count / min_len if min_len > 0 else 0.0

def should_merge_cross_page_table(
    table1: List, 
    table2: List, 
    initial_header: List,
    page_num: int,
    start_page: int,
    is_consecutive_page: bool = True
) -> Tuple[bool, str]:
    """
    智能判断是否应该合并跨页表格（通用方法）
    综合考虑多个因素，不严格限制列数差异
    
    返回: (是否合并, 原因说明)
    """
    if not table1 or not table2 or not initial_header:
        return False, "缺少必要数据"
    
    if len(table1) == 0 or len(table2) == 0:
        return False, "表格为空"
    
    # 获取列数
    cols1 = len([c for c in initial_header if c]) if initial_header else 0
    table2_first_row = table2[0] if table2 else []
    cols2 = len([c for c in table2_first_row if c]) if table2_first_row else 0
    
    if cols1 == 0 or cols2 == 0:
        return False, "列数为0"
    
    # 检查table2第一行是否是表头
    table2_has_header = has_header_row(table2_first_row)
    
    # 因素1: 无表头情况（很可能是同一表格的延续）
    if not table2_has_header:
        # 对于无表头的情况，使用更宽松的列数匹配
        # 允许±3列的差异（考虑表格在跨页时可能被拆分或格式不同）
        col_diff = abs(cols1 - cols2)
        
        # 在连续页面中，更宽松的匹配
        if is_consecutive_page and col_diff <= 3:
            # 检查数据行的结构相似性
            if len(table1) > 1 and len(table2) > 0:
                last_row1 = table1[-1]
                first_row2 = table2[0]
                
                non_empty1 = sum(1 for c in last_row1 if c and str(c).strip())
                non_empty2 = sum(1 for c in first_row2 if c and str(c).strip())
                
                # 如果非空单元格数量相似（允许±3的差异）
                if abs(non_empty1 - non_empty2) <= 4:
                    return True, f"无表头，列数差异{col_diff}列，数据行结构相似"
            
            # 对于连续页面，即使非空单元格数量不相似，只要列数差异≤3，也可能是同一表格
            # 因为表格可能在跨页时被拆分
            if col_diff <= 3:
                return True, f"无表头，连续页面，列数差异{col_diff}列，允许合并"
        
        # 即使列数差异较大，如果数据行格式非常相似，也可能合并
        if len(table1) > 1 and len(table2) > 0:
            last_row1 = table1[-1]
            first_row2 = table2[0]
            
            # 检查前几列的内容是否相似（可能是表格的一部分列）
            min_cols = min(cols1, cols2, 5)  # 至少检查前5列
            match_count = 0
            for i in range(min_cols):
                cell1 = str(last_row1[i]).strip() if i < len(last_row1) and last_row1[i] else ""
                cell2 = str(first_row2[i]).strip() if i < len(first_row2) and first_row2[i] else ""
                
                # 如果都是非空的，或者格式相似（都是数字或都是文本）
                if cell1 and cell2:
                    # 简单的内容相似性检查
                    if cell1 == cell2 or (len(cell1) > 3 and cell1[:3] == cell2[:3]):
                        match_count += 1
                    # 或者都是数字/都是文本
                    try:
                        float(cell1.replace(',', ''))
                        float(cell2.replace(',', ''))
                        match_count += 0.5  # 都是数字，部分匹配
                    except:
                        pass
            
            if match_count >= min_cols * 0.4:  # 至少40%的列匹配
                return True, f"无表头，数据行前{min_cols}列相似度{match_count/min_cols:.2f}"
    
    # 因素2: 有表头情况 - 检查表头相似度
    if table2_has_header:
        header_similarity = calculate_header_similarity(initial_header, table2_first_row)
        col_diff = abs(cols1 - cols2)
        
        # 如果表头相似度很高（≥70%），即使列数不同也可能合并
        if header_similarity >= 0.7:
            # 对于高相似度表头，允许更大的列数容差
            # 如果列数差异在合理范围内（≤5列），且表头相似度高
            if col_diff <= 5 or header_similarity >= 0.85:
                return True, f"表头相似度{header_similarity:.2f}，列数差异{col_diff}列"
        
        # 如果列数相同或接近（±3列），且表头有一定相似度（≥50%）
        # 特别是在连续页面中，降低相似度要求
        if is_consecutive_page:
            if col_diff <= 3 and header_similarity >= 0.5:
                return True, f"连续页面，列数差异{col_diff}列，表头相似度{header_similarity:.2f}"
            # 对于连续页面，即使相似度较低，如果列数差异≤3，也考虑合并
            if col_diff <= 3 and header_similarity >= 0.15:  # 降低阈值到15%
                # 额外检查数据行结构相似性
                if len(table1) > 1 and len(table2) > 1:
                    # 检查前几行数据的格式相似性
                    sample_rows = min(2, len(table1) - 1, len(table2) - 1)
                    format_match = 0
                    for i in range(1, sample_rows + 1):
                        row1 = table1[i] if i < len(table1) else []
                        row2 = table2[i] if i < len(table2) else []
                        non_empty1 = sum(1 for c in row1 if c and str(c).strip())
                        non_empty2 = sum(1 for c in row2 if c and str(c).strip())
                        if abs(non_empty1 - non_empty2) <= 3:
                            format_match += 1
                    if format_match >= sample_rows * 0.5:
                        return True, f"连续页面，列数差异{col_diff}列，表头相似度{header_similarity:.2f}，数据格式相似"
        
        # 如果列数相同或接近（±2列），且表头有一定相似度（≥60%）
        if abs(cols1 - cols2) <= 2 and header_similarity >= 0.6:
            return True, f"列数差异≤2，表头相似度{header_similarity:.2f}"
    
    # 因素3: 列数匹配（放宽到±3列，特别是在连续页面中）
    col_diff = abs(cols1 - cols2)
    if is_consecutive_page:
        if col_diff <= 3:
            # 在连续页面中，即使表头不匹配，如果列数接近，也可能是同一表格
            # 但需要检查数据行的结构相似性
            if len(table1) > 1 and len(table2) > 0:
                # 检查前几行数据的格式
                sample_rows = min(2, len(table1) - 1, len(table2))
                format_match = 0
                start_idx2 = 0 if not table2_has_header else 1
                for i in range(1, sample_rows + 1):
                    row1 = table1[i] if i < len(table1) else []
                    row2_idx = start_idx2 + i - 1
                    row2 = table2[row2_idx] if row2_idx < len(table2) else []
                    
                    non_empty1 = sum(1 for c in row1 if c and str(c).strip())
                    non_empty2 = sum(1 for c in row2 if c and str(c).strip())
                    
                    if abs(non_empty1 - non_empty2) <= 3:  # 放宽到±3
                        format_match += 1
                
                if format_match >= sample_rows * 0.5:  # 至少50%的行格式匹配
                    return True, f"连续页面，列数差异{col_diff}列，数据格式相似"
            
            # 对于连续页面，如果列数差异≤3，即使数据格式检查不完全匹配，也允许合并
            # 因为表格可能在跨页时被拆分，导致格式略有不同
            # 但需要至少有一些数据行的结构相似性
            if len(table1) > 1 and len(table2) > 0:
                last_row1 = table1[-1]
                first_data_row2 = table2[0] if not table2_has_header else (table2[1] if len(table2) > 1 else table2[0])
                
                non_empty1 = sum(1 for c in last_row1 if c and str(c).strip())
                non_empty2 = sum(1 for c in first_data_row2 if c and str(c).strip())
                
                # 如果非空单元格数量在合理范围内（至少有一半的相似性）
                if non_empty1 > 0 and non_empty2 > 0:
                    # 检查是否有重叠的非空位置
                    non_empty_positions1 = [i for i, c in enumerate(last_row1) if c and str(c).strip()]
                    non_empty_positions2 = [i for i, c in enumerate(first_data_row2) if c and str(c).strip()]
                    
                    if non_empty_positions1 and non_empty_positions2:
                        # 至少检查前min(cols1, cols2)列
                        check_cols = min(len(non_empty_positions1), len(non_empty_positions2), min(cols1, cols2))
                        overlap = len(set(non_empty_positions1[:check_cols]) & set(non_empty_positions2[:check_cols]))
                        
                        if overlap > 0:  # 只要有重叠就允许合并
                            return True, f"连续页面，列数差异{col_diff}列，非空单元格位置有重叠"
    
    # 因素4: 完全相同的列数，但表头不完全匹配（可能是表格格式问题）
    if cols1 == cols2 and cols1 >= 5:  # 至少5列，避免误判小表格
        # 检查数据行的内容格式
        if len(table1) > 1 and len(table2) > 0:
            last_row1 = table1[-1]
            first_data_row2 = table2[1] if table2_has_header and len(table2) > 1 else table2[0]
            
            # 检查非空单元格的数量和位置
            non_empty_positions1 = [i for i, c in enumerate(last_row1) if c and str(c).strip()]
            non_empty_positions2 = [i for i, c in enumerate(first_data_row2) if c and str(c).strip()]
            
            # 如果非空单元格的位置有重叠（至少30%重叠）
            if non_empty_positions1 and non_empty_positions2:
                overlap = len(set(non_empty_positions1) & set(non_empty_positions2))
                overlap_ratio = overlap / max(len(non_empty_positions1), len(non_empty_positions2))
                
                if overlap_ratio >= 0.3:
                    return True, f"列数相同，非空单元格位置重叠度{overlap_ratio:.2f}"
    
    return False, f"不满足合并条件（列数差异{col_diff}列，表头相似度{calculate_header_similarity(initial_header, table2_first_row) if table2_has_header else 0:.2f}）"

def merge_tables_with_header(table1: List, table2: List, target_header: List) -> List:
    """合并表格，使用指定的表头（确保所有跨页表格使用同一表头）"""
    if not table1:
        return table2
    if not table2:
        return table1
    if not target_header:
        return merge_tables(table1, table2)  # 如果没有指定表头，使用原有逻辑
    
    # 使用指定的表头
    merged = [target_header]
    
    # 添加table1的所有数据行（跳过表头）
    for row in table1[1:]:
        if row:
            # 更宽松的判断：只要行存在，就保留（避免遗漏数据）
            # 检查是否有任何非空单元格
            has_content = any(cell for cell in row if cell and str(cell) and str(cell).strip())
            # 或者行中有足够的单元格（可能是格式问题导致看起来为空）
            if has_content or len([c for c in row if c is not None]) > 0:
                # 确保列数一致
                while len(row) < len(target_header):
                    row.append(None)
                merged.append(row[:len(target_header)])
    
    # 检查table2的第一行是否是表头
    table2_has_header = has_header_row(table2[0]) if table2 else False
    
    # 添加table2的数据行
    start_idx = 1 if table2_has_header else 0  # 如果有表头，跳过第一行
    for row in table2[start_idx:]:
        if row:
            # 更宽松的判断：只要行存在，就保留（避免遗漏数据）
            has_content = any(cell for cell in row if cell and str(cell) and str(cell).strip())
            # 或者行中有足够的单元格
            if has_content or len([c for c in row if c is not None]) > 0:
                # 确保列数一致
                while len(row) < len(target_header):
                    row.append(None)
                merged.append(row[:len(target_header)])
    
    return merged

def merge_tables(table1: List, table2: List) -> List:
    """合并两个表格（用于跨页表格）- 增强版，处理表头重复和无表头情况"""
    if not table1:
        return table2
    if not table2:
        return table1
    
    # 检查table2是否有表头
    table2_has_header = False
    if len(table2) > 0:
        table2_has_header = has_header_row(table2[0])
    
    # 如果table2没有表头，直接合并所有行（是table1的延续）
    if not table2_has_header:
        merged = [table1[0]]  # 保留table1的表头
        
        # 添加table1的所有数据行（跳过表头）
        for row in table1[1:]:
            if row:
                # 更宽松的判断：只要行存在，就保留
                has_content = any(cell for cell in row if cell and str(cell) and str(cell).strip())
                if has_content or len([c for c in row if c is not None]) > 0:
                    merged.append(row)
        
        # 添加table2的所有行（没有表头，全部是数据）
        for row in table2:
            if row:
                # 更宽松的判断：只要行存在，就保留
                has_content = any(cell for cell in row if cell and str(cell) and str(cell).strip())
                if has_content or len([c for c in row if c is not None]) > 0:
                    # 确保列数一致
                    while len(row) < len(merged[0]):
                        row.append(None)
                    merged.append(row[:len(merged[0])])
        
        return merged
    
    # 如果table2有表头，检查是否是表头重复
    is_header_repeat = False
    if len(table1) > 0 and len(table2) > 0:
        header1 = [str(cell).strip() if cell else "" for cell in table1[0]]
        header2 = [str(cell).strip() if cell else "" for cell in table2[0]]
        
        # 如果表头相似度很高，说明是表头重复
        if len(header1) > 0 and len(header2) > 0:
            min_len = min(len(header1), len(header2))
            match_count = sum(1 for h1, h2 in zip(header1[:min_len], header2[:min_len])
                             if h1 and h2 and (h1 == h2 or h1 in h2 or h2 in h1))
            similarity = match_count / min_len if min_len > 0 else 0
            is_header_repeat = similarity >= 0.7
    
    # 合并：保留第一个表格的表头
    merged = [table1[0]]  # 表头
    
    # 添加第一个表格的数据行（跳过表头）
    for row in table1[1:]:
        if row and any(cell for cell in row if cell and str(cell).strip()):  # 确保不是空行
            merged.append(row)
    
    # 添加第二个表格的数据行
    start_idx = 1 if is_header_repeat else 0  # 如果表头重复，跳过第一行
    for row in table2[start_idx:]:
        if row and any(cell for cell in row if cell and str(cell).strip()):  # 确保不是空行
            # 确保列数一致
            while len(row) < len(merged[0]):
                row.append(None)
            merged.append(row[:len(merged[0])])
    
    return merged

def find_table_after_title(page, title_text: str, title_position: float = None) -> Optional[Tuple[List, Dict]]:
    """在页面中找到标题后的第一个表格 - 增强版"""
    try:
        # 获取表格对象（包含位置信息）
        table_objects = page.find_tables()
        
        # 提取页面中的所有表格
        tables = page.extract_tables()
        if not tables and not table_objects:
            return None
        
        # 优先使用表格对象（包含位置信息）
        if table_objects:
            best_table = None
            best_table_obj = None
            min_distance = float('inf')
            
            for table_obj in table_objects:
                table_bbox = table_obj.bbox
                table_top = table_bbox[1]  # Y坐标（从上到下）
                
                # 如果提供了标题位置，找到标题下方最近的表格
                if title_position is not None:
                    if table_top > title_position:
                        distance = table_top - title_position
                        if distance < min_distance:
                            min_distance = distance
                            best_table_obj = table_obj
                else:
                    # 如果没有位置信息，使用第一个表格
                    if best_table_obj is None:
                        best_table_obj = table_obj
            
            if best_table_obj:
                table = best_table_obj.extract()
                if table and len(table) > 1:
                    return (table, {'bbox': best_table_obj.bbox, 'object': best_table_obj})
        
        # 如果没有表格对象，使用提取的表格
        if tables:
            # 尝试找到最大的表格（通常是最相关的）
            largest_table = max(tables, key=lambda t: len(t) if t else 0)
            if largest_table and len(largest_table) > 1:
                return (largest_table, {})
        
        return None
    except Exception as e:
        print(f"  查找表格时出错: {str(e)}")
        return None

def extract_specific_content_from_pdf(pdf_path: str) -> Dict:
    """从PDF中提取指定的表格内容，只提取表格名称后的第一个表格"""
    results = {}
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            print(f"开始处理PDF，共 {total_pages} 页")
            
            # 首先扫描所有页面，找到每个表格名称的位置
            section_locations = {}
            for page_num, page in enumerate(pdf.pages, start=1):
                text = extract_text_from_page(page)
                if text:
                    for section_name, section_type in TARGET_SECTIONS:
                        if find_section_in_text(text, section_name):
                            # 找到标题在文本中的位置
                            title_index = text.find(section_name)
                            if title_index != -1:
                                # 计算标题在页面中的大概Y坐标（从上到下）
                                # 这是一个近似值，用于判断表格是否在标题下方
                                lines_before_title = len(text[:title_index].split('\n'))
                                estimated_y = lines_before_title * 15  # 假设每行约15像素
                                
                                if section_name not in section_locations:
                                    section_locations[section_name] = {
                                        'start_page': page_num,
                                        'type': section_type,
                                        'title_y': estimated_y,
                                        'end_page': page_num
                                    }
                                else:
                                    # 更新结束页面（表格可能跨多页）
                                    section_locations[section_name]['end_page'] = page_num
            
            print(f"找到 {len(section_locations)} 个表格章节")
            
            # 提取每个表格的内容
            for section_name, section_type in TARGET_SECTIONS:
                if section_name not in section_locations:
                    print(f"未找到表格: {section_name}")
                    results[section_name] = {
                        'type': section_type,
                        'tables': [],
                        'found': False
                    }
                    continue
                
                loc = section_locations[section_name]
                start_page = loc['start_page']
                title_y = loc.get('title_y', 0)
                
                # 智能确定搜索范围
                # 对于特种设备一览表，已知在121-131页，设置合理的搜索范围
                if section_name == "特种设备一览表":
                    # 从121页开始，搜索到135页（留一些缓冲，但不会太远）
                    end_page = min(start_page + 15, total_pages)
                    print(f"提取表格: {section_name} (第{start_page}页开始，智能搜索到第{end_page}页)")
                else:
                    # 对于其他表格，扩大搜索范围（最多查找后续20页）
                    end_page = min(loc.get('end_page', start_page) + 20, total_pages)
                    print(f"提取表格: {section_name} (第{start_page}-{end_page}页)")
                
                section_content = {
                    'type': section_type,
                    'tables': [],
                    'found': True
                }
                
                # 在标题所在页面查找标题后的第一个表格
                found_table = False
                page = pdf.pages[start_page - 1]
                table_result = find_table_after_title(page, section_name, title_y)
                
                if table_result:
                    table_data, table_info = table_result
                    if table_data and len(table_data) > 1:
                        cleaned_table = clean_table_data(table_data)
                        if cleaned_table:
                            section_content['tables'].append({
                                'page': start_page,
                                'data': cleaned_table,
                                'structure': table_info
                            })
                            found_table = True
                
                # 增强的跨页表格识别和合并
                # 表头只出现在表格章节后面（第一个表格），跨页表格归为最近表头
                current_table_data = section_content['tables'][0]['data'] if section_content['tables'] else None
                last_merged_page = start_page if found_table else None
                
                # 记录初始表格的表头（用于后续跨页表格）
                initial_header = None
                if current_table_data and len(current_table_data) > 0:
                    initial_header = current_table_data[0]  # 保存第一个表格的表头
                    print(f"  初始表格表头已保存: {len([c for c in initial_header if c])}列")
                
                # 连续扫描后续页面，查找跨页表格
                for page_num in range(start_page + 1, end_page + 1):
                    page = pdf.pages[page_num - 1]
                    
                    # 提取页面中的所有表格
                    tables = extract_table_improved(page)
                    
                    # 如果没有找到表格，尝试使用更宽松的策略
                    if not tables:
                        try:
                            # 尝试使用不同的策略提取
                            tables = page.extract_tables(table_settings={
                                "vertical_strategy": "lines",
                                "horizontal_strategy": "lines",
                                "snap_tolerance": 10,
                                "join_tolerance": 5,
                            })
                        except:
                            pass
                    
                    # 调试信息：显示每页找到的表格数量
                    if tables:
                        print(f"  第{page_num}页找到 {len(tables)} 个表格")
                    else:
                        print(f"  第{page_num}页未找到表格")
                    
                    merged_this_page = False
                    
                    for table in tables:
                        if table and len(table) > 1:
                            cleaned_table = clean_table_data(table)
                            
                            if not cleaned_table or len(cleaned_table) < 2:
                                continue
                            
                            # 检查是否是同一表格的延续 - 使用智能判断方法
                            if current_table_data and initial_header:
                                # 获取列数信息（用于调试）
                                cols1 = len([c for c in initial_header if c]) if initial_header else 0
                                table2_first_row = cleaned_table[0] if cleaned_table else []
                                cols2 = len([c for c in table2_first_row if c]) if table2_first_row else 0
                                
                                # 调试信息：显示所有找到的表格的列数
                                print(f"  第{page_num}页表格列数: {cols1}列 vs {cols2}列 (差异: {abs(cols1 - cols2)})")
                                
                                # 判断是否是连续页面
                                is_consecutive = (page_num == last_merged_page + 1) if last_merged_page else (page_num == start_page + 1)
                                
                                # 使用智能判断函数决定是否合并
                                should_merge, merge_reason = should_merge_cross_page_table(
                                    current_table_data,
                                    cleaned_table,
                                    initial_header,
                                    page_num,
                                    start_page,
                                    is_consecutive
                                )
                                
                                if should_merge:
                                    print(f"    -> {merge_reason}，执行合并")
                                    
                                    # 执行合并
                                    current_table_data = merge_tables_with_header(current_table_data, cleaned_table, initial_header)
                                    section_content['tables'][0]['data'] = current_table_data
                                    section_content['tables'][0]['end_page'] = page_num
                                    last_merged_page = page_num
                                    merged_this_page = True
                                    print(f"  合并跨页表格到第{page_num}页 ({merge_reason}，当前行数: {len(current_table_data)})")
                                    break
                                else:
                                    print(f"    -> {merge_reason}，跳过合并")
                            
                            elif not found_table:
                                # 如果第一页没找到，使用后续页面的第一个表格
                                section_content['tables'].append({
                                    'page': page_num,
                                    'data': cleaned_table
                                })
                                found_table = True
                                current_table_data = cleaned_table
                                initial_header = cleaned_table[0] if cleaned_table else None
                                last_merged_page = page_num
                                break
                    
                    # 智能停止搜索逻辑
                    if current_table_data and last_merged_page:
                        pages_since_last_merge = page_num - last_merged_page
                        
                        # 对于特种设备一览表，在131页后严格停止
                        if section_name == "特种设备一览表":
                            if page_num > 131:  # 超过已知范围（121-131页）
                                print(f"  超过131页，停止搜索（表格应在121-131页）")
                                break
                            # 在131页内，检查表格是否真的结束了
                            if is_table_ended(page, section_name):
                                print(f"  第{page_num}页检测到表格结束标记，停止搜索")
                                break
                            # 在131页内，如果连续2页未找到匹配表格，停止搜索
                            if pages_since_last_merge >= 2 and not merged_this_page:
                                print(f"  连续{pages_since_last_merge}页未找到匹配表格，停止搜索")
                                break
                        else:
                            # 其他表格：检查表格是否真的结束了
                            if is_table_ended(page, section_name):
                                print(f"  第{page_num}页检测到表格结束标记，停止搜索")
                                break
                            # 其他表格：连续2页未找到匹配表格，停止搜索（更严格，避免提取多余内容）
                            if pages_since_last_merge >= 2 and not merged_this_page:
                                print(f"  连续{pages_since_last_merge}页未找到匹配表格，停止搜索（避免提取多余内容）")
                                break
                    
                    # 如果找到了表格但当前页没有合并，检查是否应该停止（避免提取多余内容）
                    if found_table and current_table_data and not merged_this_page:
                        # 如果连续2页未合并，且检测到表格结束标记，停止搜索
                        if last_merged_page and (page_num - last_merged_page) >= 2:
                            if is_table_ended(page, section_name):
                                print(f"  连续2页未合并且检测到表格结束，停止搜索（避免提取多余内容）")
                                break
                        # 否则继续检查下一页（可能是表格中间有空页）
                        continue
                
                # 验证表格完整性并清理多余内容
                if section_content['tables']:
                    final_table = section_content['tables'][0]
                    if final_table.get('data'):
                        # 清理破坏连贯性的多余单元格
                        print(f"  清理前: {len(final_table['data'])}行")
                        cleaned_data = clean_duplicate_cells(final_table['data'])
                        final_table['data'] = cleaned_data
                        print(f"  清理后: {len(cleaned_data)}行")
                        
                        row_count = len(cleaned_data)
                        col_count = len(cleaned_data[0]) if cleaned_data else 0
                        print(f"  表格提取完成: {row_count}行 x {col_count}列")
                        
                        # 检查表格是否可能被截断（最后一行是否完整）
                        if row_count > 1:
                            last_row = cleaned_data[-1]
                            non_empty_cells = sum(1 for cell in last_row if cell and str(cell).strip())
                            if non_empty_cells < col_count * 0.3:  # 如果最后一行少于30%的单元格有内容
                                print(f"  警告: 表格最后一行可能不完整，建议检查")
                
                if not section_content['tables']:
                    print(f"  警告: 在指定范围内未找到表格")
                    section_content['found'] = False
                else:
                    print(f"  找到 {len(section_content['tables'])} 个表格")
                
                results[section_name] = section_content
                
                # 每处理一个章节输出进度
                processed = len([r for r in results.values() if r.get('found', False)])
                print(f"进度: {processed}/{len(TARGET_SECTIONS)} 表格已处理")
    
    except Exception as e:
        raise Exception(f"提取PDF内容时出错: {str(e)}")
    
    return results

def save_content_to_excel(content_results: Dict, output_path: str) -> Dict:
    """将提取的内容保存到Excel文件"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.utils import get_column_letter
    
    wb = Workbook()
    wb.remove(wb.active)  # 删除默认工作表
    
    stats = {
        'total_sections': 0,
        'found_sections': 0,
        'total_tables': 0,
        'total_text_sections': 0
    }
    
    for section_name, section_data in content_results.items():
        if not section_data.get('found', False):
            continue
        
        stats['found_sections'] += 1
        section_type = section_data.get('type', 'table')
        
        # 创建工作表，名称限制为31个字符
        sheet_name = section_name[:31] if len(section_name) <= 31 else section_name[:28] + "..."
        ws = wb.create_sheet(title=sheet_name)
        
        row = 1
        
        # 添加表格标题
        title_cell = ws.cell(row=row, column=1, value=section_name)
        title_cell.font = Font(bold=True, size=14, color="FFFFFF")
        title_cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        title_cell.alignment = Alignment(horizontal='center', vertical='center')
        # 合并标题行（根据表格列数动态调整）
        max_cols = 0
        for table_info in section_data.get('tables', []):
            if table_info.get('data'):
                max_cols = max(max_cols, len(table_info['data'][0]) if table_info['data'] else 0)
        if max_cols == 0:
            max_cols = 10  # 默认10列
        ws.merge_cells(f'A{row}:{get_column_letter(min(max_cols, 26))}{row}')
        row += 2
        
        # 添加表格（确保完整，不拆分）
        tables = section_data.get('tables', [])
        if tables:
            for table_idx, table_info in enumerate(tables, start=1):
                table_data = table_info.get('data', [])
                if not table_data or len(table_data) < 2:  # 至少要有表头和数据行
                    continue
                
                stats['total_tables'] += 1
                
                # 显示表格范围信息（如果是跨页表格）
                page_info = f"第{table_info.get('page', '?')}页"
                if 'end_page' in table_info and table_info['end_page'] != table_info.get('page'):
                    page_info = f"第{table_info.get('page')}-{table_info['end_page']}页"
                
                if len(tables) > 1:
                    table_title = ws.cell(row=row, column=1, value=f"表格 {table_idx} ({page_info})")
                    table_title.font = Font(bold=True, size=11, color="666666")
                    row += 1
                
                # 不检测合并单元格（用户要求：相同的字段不要合并，保持独立）
                merged_cells = []  # 禁用合并单元格功能
                
                # 写入表格数据（完整表格，不拆分）
                start_row = row
                for r_idx, table_row in enumerate(table_data, start=row):
                    # 确保列数一致（补齐缺失的列）
                    while len(table_row) < len(table_data[0]):
                        table_row.append(None)
                    
                    # 计算当前行的最大行高（基于单元格中的换行数）
                    max_lines_in_row = 1
                    
                    for c_idx, cell_value in enumerate(table_row[:len(table_data[0])], start=1):
                        # 确保单元格值中的换行符被保留
                        if cell_value is not None:
                            cell_str = str(cell_value)
                            # 计算换行数，用于设置行高
                            line_count = cell_str.count('\n') + 1
                            max_lines_in_row = max(max_lines_in_row, line_count)
                        else:
                            cell_str = None
                        
                        cell = ws.cell(row=r_idx, column=c_idx, value=cell_str)
                        # 第一行作为表头，加粗
                        if r_idx == row:
                            cell.font = Font(bold=True, size=11)
                            cell.fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
                            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                        else:
                            cell.alignment = Alignment(vertical='top', horizontal='left', wrap_text=True)
                    
                    # 根据单元格中的换行数设置行高（每行约15像素，最小20像素）
                    if r_idx != row:  # 数据行
                        row_height = max(20, max_lines_in_row * 15)
                        ws.row_dimensions[r_idx].height = min(row_height, 400)  # 最大400像素
                    else:  # 表头行
                        ws.row_dimensions[r_idx].height = max(25, max_lines_in_row * 15)
                
                # 应用单元格合并（保持原表格格式）
                for merge in merged_cells:
                    try:
                        # merge格式: (start_row, start_col, end_row, end_col)
                        # Excel格式: 行和列从1开始，需要加上start_row偏移
                        excel_start_row = merge[0] + start_row
                        excel_start_col = merge[1] + 1
                        excel_end_row = merge[2] + start_row
                        excel_end_col = merge[3] + 1
                        
                        # 只合并有效的单元格（至少2x2）
                        if excel_end_row > excel_start_row or excel_end_col > excel_start_col:
                            merge_range = f"{get_column_letter(excel_start_col)}{excel_start_row}:{get_column_letter(excel_end_col)}{excel_end_row}"
                            ws.merge_cells(merge_range)
                            # 设置合并后单元格的对齐方式
                            merged_cell = ws.cell(row=excel_start_row, column=excel_start_col)
                            merged_cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    except Exception as e:
                        # 如果合并失败，继续处理其他单元格
                        pass
                
                # 自动调整列宽（确保完整显示）
                for col in range(1, len(table_data[0]) + 1):
                    col_letter = get_column_letter(col)
                    max_length = 10  # 最小宽度
                    for row_num in range(row, row + len(table_data)):
                        cell_value = ws.cell(row=row_num, column=col).value
                        if cell_value:
                            # 计算单元格内容长度（考虑换行）
                            cell_str = str(cell_value)
                            lines = cell_str.split('\n')
                            max_length = max(max_length, max(len(line) for line in lines) if lines else len(cell_str))
                    # 设置列宽，最大不超过60
                    ws.column_dimensions[col_letter].width = min(max_length + 2, 60)
                
                row += len(table_data) + 2  # 表格后留空行
    
    stats['total_sections'] = len(TARGET_SECTIONS)
    
    # 如果没有找到任何内容，创建一个说明工作表
    if stats['found_sections'] == 0:
        ws = wb.create_sheet(title="说明")
        ws.cell(row=1, column=1, value="未在PDF中找到指定的章节内容")
    
    wb.save(output_path)
    return stats


def save_content_to_docx(content_results: Dict, output_path: str) -> Dict:
    """
    将提取的表格内容保存为高保真 Word 文档（.docx）。
    逻辑参考 save_content_to_excel：每个章节/表格块对应一个 Word 表格。
    表格样式：Table Grid、表头浅灰加粗、等线/宋体 9 磅、居中对齐、宽度 100%。
    """
    from docx import Document
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    # 全局默认字体：等线 9 磅（Normal 样式）
    style = doc.styles['Normal']
    style.font.name = '等线'
    style.font.size = Pt(9)
    try:
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
    except Exception:
        pass

    stats = {
        'total_sections': 0,
        'found_sections': 0,
        'total_tables': 0,
    }

    for section_name, section_data in content_results.items():
        if not section_data.get('found', False):
            continue
        stats['found_sections'] += 1
        tables = section_data.get('tables', [])
        if not tables:
            continue

        for table_info in tables:
            table_data = table_info.get('data', [])
            if not table_data or len(table_data) < 1:
                continue
            stats['total_tables'] += 1

            # 添加表名（表格标题），放在表格上方
            title_para = doc.add_paragraph()
            run = title_para.add_run(section_name)
            run.bold = True
            run.font.name = '等线'
            run.font.size = Pt(10.5)
            try:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
            except Exception:
                pass
            title_para.paragraph_format.space_after = Pt(6)

            # 确定行数列数（以第一行长度为列数，不足的列补空）
            num_rows = len(table_data)
            num_cols = max(len(row) for row in table_data) if table_data else 0
            if num_cols == 0:
                continue
            # 统一每行列数
            for row in table_data:
                while len(row) < num_cols:
                    row.append('')

            table = doc.add_table(rows=num_rows, cols=num_cols, style='Table Grid')
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 表格宽度设为页面可用宽度 100%
            try:
                section = doc.sections[0]
                table_width = section.page_width - section.left_margin - section.right_margin
                table.width = table_width
            except Exception:
                pass

            for r_idx, row_data in enumerate(table_data):
                row = table.rows[r_idx]
                for c_idx, value in enumerate(row_data[:num_cols]):
                    cell = row.cells[c_idx]
                    # 文本还原，保留换行符 \n
                    cell.text = str(value) if value is not None else ''
                    # 单元格内容垂直居中
                    tcPr = cell._tc.get_or_add_tcPr()
                    vAlign = OxmlElement('w:vAlign')
                    vAlign.set(qn('w:val'), 'center')
                    tcPr.append(vAlign)
                    # 字体与大小：等线/宋体 9 磅
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = '等线'
                            run.font.size = Pt(9)
                            try:
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
                            except Exception:
                                pass
                    # 第一行作为表头：背景 D3D3D3，加粗
                    if r_idx == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:fill'), 'D3D3D3')
                        cell._tc.get_or_add_tcPr().append(shd)

            # 表格后留空
            doc.add_paragraph()

    stats['total_sections'] = len(content_results)
    if stats['found_sections'] == 0:
        doc.add_paragraph('未在PDF中找到指定的章节内容。')
    doc.save(output_path)
    return stats


def _clean_cell_text(text: str) -> str:
    """清洗单元格文本：仅保留 \\n 换行，去除 \\r、\\a 等不可见字符。"""
    if not text or not isinstance(text, str):
        return ""
    s = text.replace("\r\n", "\n").replace("\r", "\n").replace("\a", "\n")
    s = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", s)
    return s.strip()


def _normalize_docx_table_title(paragraph_text: Optional[str]) -> Optional[str]:
    """
    仅认可两种表格名称格式，用于 Word 输入时的表名提取：
    （1）评价报告摘要
    （2）表 x-y 文字（如：表 2-2 企业现有已投产项目各车间产品分布情况表）
    若段落文本不符合这两种之一，返回 None（调用方用「表格 N」等占位）。
    """
    if not paragraph_text or not isinstance(paragraph_text, str):
        return None
    s = paragraph_text.strip()
    if not s:
        return None
    if "评价报告摘要" in s:
        return "评价报告摘要"
    # 表 x-y 文字：表 + 数字 + 横线(-/－/–) + 数字 + 空格 + 文字（如：表 2-2 企业...）
    m = re.match(r"^表\s*\d+\s*[-－–]\s*\d+\s+.+", s)
    if m:
        return s
    m = re.search(r"表\s*\d+\s*[-－–]\s*\d+\s+.+", s)
    if m:
        return m.group(0).strip()
    return None


def _clean_docx_title_line(paragraph_text: Optional[str]) -> Optional[str]:
    """从段落文本中提取“上一行表名”：清洗不可见字符，仅取第一条非空行。"""
    if not paragraph_text or not isinstance(paragraph_text, str):
        return None
    s = _clean_cell_text(paragraph_text)
    if not s:
        return None
    for line in s.split("\n"):
        t = re.sub(r"\s+", " ", (line or "").strip())
        if t:
            return t
    return None


def _dedupe_title_text(title: str) -> str:
    """去掉表名中重复内容（整段重复/连续重复 token），用于前端展示。"""
    if not title or not isinstance(title, str):
        return ""
    s = re.sub(r"\s+", " ", title).strip()
    if not s:
        return ""
    # 整段重复（如 “X X” 或 “X   X”）
    m = re.match(r"^(.+?)(?:\s+\1)+$", s)
    if m:
        s = m.group(1).strip()
    # token 级去重（连续重复）
    tokens = s.split(" ")
    out = []
    prev = None
    for tok in tokens:
        if tok and tok != prev:
            out.append(tok)
        prev = tok
    # 半段重复（token 序列前半 == 后半）
    if len(out) % 2 == 0 and out[: len(out) // 2] == out[len(out) // 2 :]:
        out = out[: len(out) // 2]
    return " ".join(out).strip()


def _get_docx_table_groups(docx_path: str):
    """
    从 Word 正文中获取表格分组（按表名去重），并返回 table_id -> 同组所有 table_id 的映射。
    表名取每个表格正上方的一行段落文字（清洗、去重后），为空则回退为“表格N”。
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(docx_path)
    body = doc.element.body
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")

    last_para_text = None
    groups_by_name = {}  # name -> group dict
    groups_in_order = []
    table_index = 0

    for child in body:
        if child.tag == p_tag:
            try:
                t = "".join((n.text or "") for n in child.iter() if hasattr(n, "text"))
                t = (t or "").strip()
                if t:
                    last_para_text = t
            except Exception:
                pass
            continue

        if child.tag != tbl_tag:
            continue

        table_id = f"table_{table_index}"
        raw_title = _clean_docx_title_line(last_para_text)
        # 优先使用规范化标题，否则使用原始上一行（两者都做重复清理）
        norm = _normalize_docx_table_title(raw_title) if raw_title else None
        name = _dedupe_title_text(norm or raw_title or f"表格{table_index + 1}")

        if name not in groups_by_name:
            group = {
                "id": table_id,  # 代表 id（前端勾选用）
                "name": name,
                "page": table_index + 1,  # Word 无页码，用序号占位避免 undefined
                "table_num": len(groups_in_order) + 1,
                "table_ids": [table_id],
                "count": 1,
            }
            groups_by_name[name] = group
            groups_in_order.append(group)
        else:
            group = groups_by_name[name]
            group["table_ids"].append(table_id)
            group["count"] += 1

        table_index += 1

    id_to_group_ids = {}
    for group in groups_in_order:
        ids = group.get("table_ids", [])
        for tid in ids:
            id_to_group_ids[tid] = ids
        # 代表 id 也映射到整组
        id_to_group_ids[group["id"]] = ids

    return groups_in_order, id_to_group_ids


def _extract_table_grid_from_docx_table(table) -> List[List[str]]:
    """
    从 python-docx 的 Table 对象提取规整的「列表嵌套列表」矩阵。
    合并单元格：按坐标 table.cell(row_idx, col_idx) 读取，合并区域仅在首个逻辑格保留文本，其余填空，保证列不错位。
    清洗 \\r、\\a，仅保留 \\n。
    """
    try:
        nrows = len(table.rows)
        if nrows == 0:
            return []
        ncols = 0
        try:
            from docx.oxml.ns import qn
            tbl = table._tbl
            grid = getattr(tbl, "tblGrid", None) or (tbl.find(qn("w:tblGrid")) if hasattr(tbl, "find") else None)
            if grid is not None and hasattr(grid, "findall"):
                ncols = len(grid.findall(qn("w:gridCol")))
        except Exception:
            pass
        if ncols <= 0:
            for c in range(256):
                try:
                    table.cell(0, c)
                    ncols = c + 1
                except Exception:
                    break
        if ncols <= 0:
            return []

        seen_cell_ids = set()
        grid_data = []
        for r in range(nrows):
            row_data = []
            for c in range(ncols):
                try:
                    cell = table.cell(r, c)
                    cell_id = id(cell)
                    if cell_id not in seen_cell_ids:
                        seen_cell_ids.add(cell_id)
                        row_data.append(_clean_cell_text(cell.text))
                    else:
                        row_data.append("")
                except Exception:
                    row_data.append("")
            grid_data.append(row_data)
        return grid_data
    except Exception:
        return []


def extract_tables_from_docx(docx_path: str) -> List[Dict]:
    """
    从 Word (.docx) 中提取所有表格，返回与 PDF 提取结果一致的 tables_data 结构。
    每个元素: {'id': str, 'page': int, 'title': str, 'data': List[List[str]]}
    合并单元格：仅在第一个逻辑单元格保留内容，其余填空，保证矩形矩阵、列不错位。
    """
    from docx import Document
    from docx.oxml.ns import qn

    tables_data = []
    try:
        doc = Document(docx_path)
        body = doc.element.body
        last_paragraph_text = None
        table_index = 0
        for child in body:
            tag = child.tag if hasattr(child, "tag") else ""
            if "tbl" in tag or child.tag == qn("w:tbl"):
                # 仅保留两种表名：（1）评价报告摘要（2）表 x-y 文字；其他文字不作为表名
                raw_title = (last_paragraph_text and last_paragraph_text.strip()) or None
                title = _normalize_docx_table_title(raw_title) if raw_title else None
                if title and len(title) > 120:
                    title = title[:117] + "..."
                if table_index < len(doc.tables):
                    tbl = doc.tables[table_index]
                    grid = _extract_table_grid_from_docx_table(tbl)
                    if grid:
                        tables_data.append({
                            "id": f"table_{table_index}",
                            "page": table_index,
                            "title": title or f"表格{table_index + 1}",
                            "data": grid,
                        })
                    table_index += 1
            elif "p" in tag or child.tag == qn("w:p"):
                try:
                    t = "".join(n.text or "" for n in child.iter() if hasattr(n, "text"))
                    if t and t.strip():
                        last_paragraph_text = t.strip()
                except Exception:
                    pass
        # 若 body 迭代与 doc.tables 顺序不一致，则退化为仅按 doc.tables 顺序提取
        if table_index < len(doc.tables) and not tables_data:
            for i, tbl in enumerate(doc.tables):
                grid = _extract_table_grid_from_docx_table(tbl)
                if grid:
                    tables_data.append({
                        "id": f"table_{i}",
                        "page": i,
                        "title": f"表格{i + 1}",
                        "data": grid,
                    })
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise e
    return tables_data


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """处理文件上传"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': '文件名为空'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': '只支持 PDF 或 Word (.docx) 文件'}), 400
    
    try:
        # 保存上传的文件
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        unique_filename = f"{timestamp}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(filepath)
        
        return jsonify({
            'message': '文件上传成功',
            'filename': unique_filename,
            'original_filename': filename
        }), 200
    
    except Exception as e:
        return jsonify({'error': f'上传文件时出错: {str(e)}'}), 500

@app.route('/api/tables', methods=['POST'])
def get_tables_list():
    """获取 PDF 或 Word 中所有表格的列表"""
    data = request.get_json()
    
    if not data or 'filename' not in data:
        return jsonify({'error': '缺少文件名参数'}), 400
    
    filename = data['filename']
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    
    if not os.path.exists(filepath):
        return jsonify({'error': '文件不存在'}), 404
    
    try:
        ext = (filename or "").rsplit(".", 1)[-1].lower()
        is_word = (ext == "docx") or _is_docx_by_magic(filepath)
        if is_word:
            # Word：表名按“表格上一行文字”提取；重复表名只展示一次
            groups, _ = _get_docx_table_groups(filepath)
            # 前端只需要 id/name/page/table_num；table_ids/count 作为附加信息保留
            n = len(groups)
            print(f"[调试] Word 获取到 {n} 个表格分组（去重表名）")
            return jsonify({
                'message': '成功获取表格列表',
                'total_tables': n,
                'display_tables': n,
                'tables': groups
            }), 200

        # PDF：使用 extract_all_tables 模块
        import importlib.util
        from pathlib import Path
        
        extract_script_path = Path(__file__).resolve().parent.parent / 'extract_all_tables.py'
        extract_script_str = os.path.normpath(str(extract_script_path))
        
        if not os.path.exists(extract_script_str):
            return jsonify({'error': f'提取脚本不存在: {extract_script_str}'}), 500
        
        spec = importlib.util.spec_from_file_location("extract_all_tables", extract_script_str)
        if spec is None or spec.loader is None:
            return jsonify({'error': '无法创建模块规范'}), 500
        
        extract_module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(extract_module)
        
        if not hasattr(extract_module, 'get_all_tables_info'):
            return jsonify({'error': '模块中未找到 get_all_tables_info 函数'}), 500
        
        all_tables_info = extract_module.get_all_tables_info(filepath)
        print(f"[调试] 获取到 {len(all_tables_info)} 个表格")
        
        if hasattr(extract_module, 'filter_tables_for_display'):
            filtered_tables = extract_module.filter_tables_for_display(all_tables_info)
        else:
            filtered_tables = all_tables_info
        
        return jsonify({
            'message': '成功获取表格列表',
            'total_tables': len(all_tables_info),
            'display_tables': len(filtered_tables),
            'tables': filtered_tables
        }), 200
    
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"获取表格列表时出错: {str(e)}")
        print(error_trace)
        return jsonify({
            'error': f'获取表格列表时出错: {str(e)}',
            'details': error_trace if app.debug else None
        }), 500


@app.route('/api/extract', methods=['POST'])
def extract_tables():
    """提取 PDF 或 Word 中的表格，统一输出为 Word (.docx)"""
    data = request.get_json()
    
    if not data or 'filename' not in data:
        return jsonify({'error': '缺少文件名参数'}), 400
    
    filename = data['filename']
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    
    selected_table_ids = data.get('selected_table_ids', None)
    
    if not os.path.exists(filepath):
        return jsonify({'error': '文件不存在'}), 404
    
    try:
        ext = (filename or "").rsplit(".", 1)[-1].lower()
        is_word = (ext == "docx") or _is_docx_by_magic(filepath)
        if is_word:
            # ---------- Word 输入：仅删除表格外内容，表格原样保留 ----------
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"all_tables_{timestamp}.docx"
            if not output_filename.lower().endswith('.docx'):
                output_filename = output_filename.rstrip('.') + '.docx'
            final_output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            os.makedirs(os.path.dirname(final_output_path), exist_ok=True)
            try:
                # 若前端表格列表做了“表名去重”，这里需要把代表 id 展开为同名的所有表格 id
                expanded_ids = None
                if selected_table_ids:
                    _, id_to_group_ids = _get_docx_table_groups(filepath)
                    expanded = []
                    seen = set()
                    for tid in selected_table_ids:
                        for x in id_to_group_ids.get(tid, [tid]):
                            if x not in seen:
                                seen.add(x)
                                expanded.append(x)
                    expanded_ids = expanded
                kept = word_remove_non_table_content(filepath, final_output_path, expanded_ids)
            except Exception as e:
                import traceback
                traceback.print_exc()
                return jsonify({
                    'error': f'处理 Word 文档时出错: {str(e)}',
                    'details': traceback.format_exc() if app.debug else None
                }), 500
            if kept == 0:
                return jsonify({
                    'error': '文档中无表格或未选中任何表格',
                    'hint': '请确认 Word 文档中包含表格并勾选需要保留的表格'
                }), 500
            try:
                if os.path.exists(filepath):
                    os.remove(filepath)
            except Exception:
                pass
            return jsonify({
                'message': '已删除表格外内容，表格原样保留',
                'total_tables': kept,
                'total_pages': kept,
                'output_filename': output_filename,
                'download_url': f'/api/download/{output_filename}',
                'file_type': 'docx',
                'found_sections': kept,
                'total_sections': kept
            }), 200

        # ---------- PDF 输入：使用 extract_all_tables 提取，统一输出 .docx ----------
        import importlib.util
        from pathlib import Path
        import shutil
        
        extract_script_path = Path(__file__).resolve().parent.parent / 'extract_all_tables.py'
        extract_script_str = os.path.normpath(str(extract_script_path))
        
        print(f"检查提取脚本路径: {extract_script_str}")
        print(f"脚本是否存在: {os.path.exists(extract_script_str)}")
        
        if not os.path.exists(extract_script_str):
            error_msg = f'提取脚本不存在: {extract_script_str}'
            print(f"错误: {error_msg}")
            return jsonify({'error': error_msg}), 500
        
        try:
            print(f"正在加载模块: {extract_script_str}")
            spec = importlib.util.spec_from_file_location("extract_all_tables", extract_script_str)
            if spec is None or spec.loader is None:
                error_msg = f'无法创建模块规范: {extract_script_str}'
                print(f"错误: {error_msg}")
                return jsonify({'error': error_msg}), 500
            
            extract_module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(extract_module)
            print(f"模块加载成功")
            
            # 检查函数是否存在
            if not hasattr(extract_module, 'extract_all_tables_from_pdf'):
                error_msg = '模块中未找到 extract_all_tables_from_pdf 函数'
                print(f"错误: {error_msg}")
                return jsonify({'error': error_msg}), 500
            print(f"函数 extract_all_tables_from_pdf 存在")
            
        except Exception as e:
            import traceback
            error_trace = traceback.format_exc()
            error_msg = f'加载模块时出错: {str(e)}'
            print(f"错误: {error_msg}")
            print(f"详细错误:\n{error_trace}")
            return jsonify({'error': error_msg, 'details': error_trace if app.debug else None}), 500
        
        # 仅提取表格数据并生成 Word，不生成 PDF
        temp_output_dir = os.path.join(app.config['OUTPUT_FOLDER'], 'temp_extracted')
        os.makedirs(temp_output_dir, exist_ok=True)
        
        print(f"=" * 60)
        print(f"开始提取所有表格（输出 Word）: {filename}")
        print(f"文件路径: {filepath}")
        print(f"输出目录: {temp_output_dir}")
        print(f"=" * 60)
        
        print(f"准备调用 extract_all_tables_from_pdf(output_format='docx')")
        print(f"  PDF文件: {filepath}")
        if selected_table_ids:
            print(f"  选择的表格ID: {selected_table_ids}")
        else:
            print(f"  提取所有表格")
        
        try:
            result = extract_module.extract_all_tables_from_pdf(filepath, temp_output_dir, selected_table_ids, output_format='docx')
        except FileNotFoundError as e:
            # 文件不存在错误
            error_msg = f'PDF文件不存在: {str(e)}'
            print(f"错误: {error_msg}")
            return jsonify({
                'error': error_msg,
                'error_type': 'FileNotFoundError',
                'hint': '请确保PDF文件已正确上传'
            }), 500
        except Exception as e:
            # 其他异常
            import traceback
            error_trace = traceback.format_exc()
            error_msg = f'调用提取函数时出错: {str(e)}'
            print(f"=" * 60)
            print(f"错误: {error_msg}")
            print(f"错误类型: {type(e).__name__}")
            print(f"详细错误:\n{error_trace}")
            print(f"=" * 60)
            
            # 检查是否有traceback属性（我们添加的详细错误信息）
            if hasattr(e, 'traceback'):
                print(f"函数内部错误详情:\n{e.traceback}")
            
            return jsonify({
                'error': error_msg,
                'error_type': type(e).__name__,
                'details': error_trace if app.debug else '启用debug模式查看详细信息'
            }), 500
        
        # 检查结果是否有效
        if not result:
            error_msg = '提取表格失败，函数返回None（这不应该发生，因为现在函数会抛出异常）'
            print(f"错误: {error_msg}")
            print(f"这可能是因为函数内部有未捕获的异常")
            return jsonify({
                'error': error_msg,
                'hint': '请检查后端控制台的详细错误信息，函数应该抛出异常而不是返回None'
            }), 500
        
        # 验证结果格式
        if not isinstance(result, dict):
            error_msg = f'提取函数返回了意外的类型: {type(result)}，期望dict'
            print(f"错误: {error_msg}")
            return jsonify({
                'error': error_msg,
                'hint': '函数应该返回一个包含total_pages, total_tables, output_dir, tables_data的字典'
            }), 500
        
        print(f"提取函数返回结果: {type(result)}")
        print(f"结果键: {result.keys() if isinstance(result, dict) else 'N/A'}")
        
        tables_data = result.get('tables_data', [])
        if not tables_data:
            return jsonify({
                'error': '未提取到任何表格数据',
                'hint': '请确认PDF中包含表格且提取函数返回了 tables_data'
            }), 500
        
        # 将 tables_data 转为 save_content_to_docx 所需的 content_results 结构
        content_results = {}
        for t in tables_data:
            title = t.get('title') or t.get('id', '表格')
            if title not in content_results:
                content_results[title] = {'found': True, 'type': 'table', 'tables': []}
            content_results[title]['tables'].append({'data': t.get('data', [])})
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"all_tables_{timestamp}.docx"
        if not output_filename.lower().endswith('.docx'):
            output_filename = output_filename.rstrip('.') + '.docx'
        final_output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        os.makedirs(os.path.dirname(final_output_path), exist_ok=True)
        
        try:
            save_content_to_docx(content_results, final_output_path)
        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({
                'error': f'生成 Word 文档时出错: {str(e)}',
                'details': traceback.format_exc() if app.debug else None
            }), 500
        
        print(f"Word 文档已保存到: {final_output_path}")
        
        # 清理上传的文件
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
                print(f"已清理上传文件: {filepath}")
        except Exception as e:
            print(f"清理上传文件时出错: {str(e)}")
        
        print(f"=" * 60)
        print(f"提取完成: {result.get('total_tables', 0)} 个表格，已生成 .docx")
        print(f"=" * 60)
        
        return jsonify({
            'message': '表格提取成功（Word 格式）',
            'total_tables': result.get('total_tables', 0),
            'total_pages': result.get('total_pages', 0),
            'output_filename': output_filename,
            'download_url': f'/api/download/{output_filename}',
            'file_type': 'docx',
            'found_sections': result.get('total_tables', 0),
            'total_sections': result.get('total_tables', 0)
        }), 200
    
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"=" * 60)
        print(f"提取表格时出错: {str(e)}")
        print(f"错误详情:")
        print(error_trace)
        print(f"=" * 60)
        
        # 清理上传的文件
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
        except:
            pass
        
        # 返回更详细的错误信息
        error_message = str(e)
        error_type = type(e).__name__
        
        # 提取关键错误信息
        error_summary = f'提取表格时出错: {error_message}'
        
        return jsonify({
            'error': error_summary,
            'error_type': error_type,
            'error_message': error_message,
            'error_details': error_trace if app.debug else '启用debug模式查看详细信息'
        }), 500

@app.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    """下载提取结果文件（支持 PDF 与 Word .docx）"""
    # 禁止路径穿越
    if '..' in filename or '/' in filename or '\\' in filename:
        return jsonify({'error': '非法文件名'}), 400
    filepath = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    
    if not os.path.exists(filepath):
        return jsonify({'error': '文件不存在'}), 404
    
    # 根据后缀设置 mimetype（本应用提取结果统一为 .docx）
    mimetype = 'application/pdf'
    if filename.lower().endswith('.docx'):
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
    try:
        return send_file(
            filepath,
            as_attachment=True,
            download_name=filename,
            mimetype=mimetype
        )
    except Exception as e:
        return jsonify({'error': f'下载文件时出错: {str(e)}'}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    """健康检查接口"""
    return jsonify({'status': 'ok', 'message': '服务运行正常'}), 200

@app.route('/api/test-extract-module', methods=['GET'])
def test_extract_module():
    """测试extract_all_tables模块是否可以正常加载"""
    try:
        import sys
        import importlib.util
        from pathlib import Path
        
        extract_script_path = Path(__file__).resolve().parent.parent / 'extract_all_tables.py'
        extract_script_str = os.path.normpath(str(extract_script_path))
        
        result = {
            'script_path': extract_script_str,
            'script_exists': os.path.exists(extract_script_str),
            'module_loaded': False,
            'function_exists': False,
            'error': None
        }
        
        if not os.path.exists(extract_script_str):
            result['error'] = f'脚本不存在: {extract_script_str}'
            return jsonify(result), 200
        
        try:
            spec = importlib.util.spec_from_file_location("extract_all_tables", extract_script_str)
            if spec is None or spec.loader is None:
                result['error'] = '无法创建模块规范'
                return jsonify(result), 200
            
            extract_module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(extract_module)
            result['module_loaded'] = True
            
            if hasattr(extract_module, 'extract_all_tables_from_pdf'):
                result['function_exists'] = True
            else:
                result['error'] = '函数 extract_all_tables_from_pdf 不存在'
                
        except Exception as e:
            import traceback
            result['error'] = str(e)
            result['traceback'] = traceback.format_exc()
        
        return jsonify(result), 200
        
    except Exception as e:
        import traceback
        return jsonify({
            'error': str(e),
            'traceback': traceback.format_exc()
        }), 500

@app.route("/")
def index():
    """提供前端页面"""
    return send_from_directory(FRONTEND_DIR, "index.html")

if __name__ == '__main__':
    port = PORT
    debug = app.config['DEBUG']
    print("=" * 50)
    print("PDF表格提取服务启动中...")
    print("  >>> 提取结果将输出为 Word (.docx) 格式 <<<")
    print("=" * 50)
    print(f"运行目录: {os.path.dirname(os.path.abspath(__file__))}")
    print(f"前端界面: http://0.0.0.0:{port}")
    print(f"API接口: http://0.0.0.0:{port}/api")
    print("=" * 50)
    if debug:
        print("\n已注册的路由:")
        for rule in app.url_map.iter_rules():
            print(f"  {rule.rule} -> {rule.endpoint} [{', '.join(rule.methods)}]")
        print("=" * 50)
    app.run(debug=debug, host='0.0.0.0', port=port)
